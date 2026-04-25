#!/usr/bin/env python3
"""
comprehensive_demo_benchmark.py — Medha Capability Demonstration & Benchmark

Tests every document type (PDF, DOCX, TXT, MD) across all workflows:
  1. INGEST: Parse, chunk, embed each document type
  2. QUERY: Reasoning accuracy on document content
  3. RFI: Draft quality with source citations
  4. SCAN: Contradiction detection between specs and drawings

Generates a credibility report for stakeholder demonstration.

RESEARCH BASIS:
- [CITE: SaaSFactor2025] TTFV <10 min or users churn. Target: all workflows <60s total.
- [CITE: Li2024] Multi-agent consensus reduces false negatives by 80%.
- [CITE: LayerTeam2025] RFIs with citations get answered 40% faster.
- [CITE: DocumentCrunch2025] Good RFIs need exact drawing/spec references.
- [CITE: Nielsen1994] System status visibility: report per-step timing.
"""

import sys
import time
import json
import hashlib
from pathlib import Path
from datetime import datetime

sys.path.insert(0, str(Path(__file__).parent.parent / "agents"))
from orchestrator import run_workflow

# ---------------------------------------------------------------------------
# CONFIGURATION
# ---------------------------------------------------------------------------
TEST_PROJECT = "demo_benchmark_" + datetime.now().strftime("%Y%m%d_%H%M%S")
BENCHMARK_DIR = Path(__file__).resolve().parent.parent.parent / "benchmark_docs"
REPORT_DIR = Path(__file__).resolve().parent.parent.parent / "benchmark_reports"
TIMING_BUDGET_SEC = 60.0

# Expected ground-truth answers for reasoning validation
BENCHMARK_QUESTIONS = [
    {
        "question": "What is the minimum concrete compressive strength required at 28 days?",
        "expected_keywords": ["30", "MPa", "compressive", "strength"],
        "source_doc": "STRUCT_SPEC.txt",
    },
    {
        "question": "What fire rating is required for structural steel beams?",
        "expected_keywords": ["2-hour", "fire", "rating", "steel"],
        "source_doc": "FIRE_PROTECTION_SPEC.txt",
    },
    {
        "question": "What is the design outdoor air temperature for HVAC sizing?",
        "expected_keywords": ["35", "°C", "outdoor", "air"],
        "source_doc": "MECH_SPEC_HVAC.txt",
    },
    {
        "question": "What is the foundation type for the administration building?",
        "expected_keywords": ["spread", "footing", "foundation"],
        "source_doc": "ARCH_DRAWING_NOTES.txt",
    },
    {
        "question": "What is the maximum allowable floor deflection per ACI code?",
        "expected_keywords": ["L/360", "deflection", "ACI"],
        "source_doc": "STRUCT_SPEC.txt",
    },
]

CONTRADICTION_EXPECTATIONS = [
    {
        "description": "Concrete strength in STRUCT_SPEC vs ARCH_DRAWING_NOTES",
        "expected_severity": "critical",
    },
    {
        "description": "Fire rating in FIRE_PROTECTION_SPEC vs STRUCT_SPEC",
        "expected_severity": "high",
    },
]

RFI_TEST_QUESTIONS = [
    {
        "question": "The structural spec calls for 30 MPa concrete but the drawing notes reference 25 MPa. Which value governs?",
        "expected_references": ["STRUCT_SPEC", "ARCH_DRAWING_NOTES"],
    },
    {
        "question": "Does the HVAC spec require fire dampers at all penetrations, or only those exceeding 2-hour ratings?",
        "expected_references": ["MECH_SPEC_HVAC", "FIRE_PROTECTION_SPEC"],
    },
]


def log(msg: str):
    print(f"[{datetime.now().strftime('%H:%M:%S')}] {msg}")


def generate_sample_documents():
    """Generate PDF, DOCX, TXT, and MD sample documents for benchmark testing."""
    BENCHMARK_DIR.mkdir(parents=True, exist_ok=True)
    log(f"Generating sample documents in {BENCHMARK_DIR}")

    # Read existing TXT samples as source content
    txt_dir = Path(__file__).resolve().parent.parent.parent / "sample_docs"
    docs = {}
    for f in sorted(txt_dir.glob("*.txt")):
        docs[f.stem] = f.read_text()

    generated = []

    # Generate DOCX files
    try:
        from docx import Document
        for name, content in docs.items():
            docx_path = BENCHMARK_DIR / f"{name}.docx"
            doc = Document()
            doc.add_heading(name.replace("_", " "), level=1)
            for para in content.strip().split("\n\n"):
                doc.add_paragraph(para.strip())
            doc.save(str(docx_path))
            generated.append(("docx", docx_path))
    except Exception as e:
        log(f"WARNING: DOCX generation failed: {e}")

    # Generate PDF files
    try:
        from reportlab.lib.pagesizes import letter
        from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
        from reportlab.lib.styles import getSampleStyleSheet
        for name, content in docs.items():
            pdf_path = BENCHMARK_DIR / f"{name}.pdf"
            pdf = SimpleDocTemplate(str(pdf_path), pagesize=letter)
            styles = getSampleStyleSheet()
            story = [Paragraph(f"<b>{name.replace('_', ' ')}</b>", styles["Heading1"]), Spacer(1, 12)]
            for para in content.strip().split("\n\n"):
                story.append(Paragraph(para.strip(), styles["Normal"]))
                story.append(Spacer(1, 6))
            pdf.build(story)
            generated.append(("pdf", pdf_path))
    except Exception as e:
        log(f"WARNING: PDF generation failed: {e}")

    # Generate MD files
    try:
        for name, content in docs.items():
            md_path = BENCHMARK_DIR / f"{name}.md"
            md_content = f"# {name.replace('_', ' ')}\n\n{content}"
            md_path.write_text(md_content)
            generated.append(("md", md_path))
    except Exception as e:
        log(f"WARNING: MD generation failed: {e}")

    # Copy original TXT files
    for name, content in docs.items():
        txt_path = BENCHMARK_DIR / f"{name}.txt"
        txt_path.write_text(content)
        generated.append(("txt", txt_path))

    log(f"Generated {len(generated)} sample documents: {len([g for g in generated if g[0]=='pdf'])} PDF, {len([g for g in generated if g[0]=='docx'])} DOCX, {len([g for g in generated if g[0]=='md'])} MD, {len([g for g in generated if g[0]=='txt'])} TXT")
    return generated


def test_ingest(doc_type: str, file_path: Path) -> dict:
    """Test ingestion of a single document."""
    name = f"Ingest {doc_type.upper()}: {file_path.name}"
    t0 = time.time()
    try:
        result = run_workflow("ingest", {
            "project": TEST_PROJECT,
            "file": str(file_path),
            "doc_type": "spec" if "SPEC" in file_path.name else "drawing",
            "use_docling": doc_type == "pdf",
        })
        duration = round(time.time() - t0, 2)
        success = "error" not in result
        log(f"{'PASS' if success else 'FAIL'}: {name} ({duration}s)")
        if not success:
            log(f"  ERROR: {result.get('error')}")
        return {
            "name": name,
            "doc_type": doc_type,
            "file": file_path.name,
            "success": success,
            "duration_sec": duration,
            "error": result.get("error") if not success else None,
        }
    except Exception as e:
        duration = round(time.time() - t0, 2)
        log(f"EXCEPTION: {name} ({duration}s) — {e}")
        return {"name": name, "doc_type": doc_type, "file": file_path.name, "success": False, "duration_sec": duration, "error": str(e)}


def test_query(question_data: dict) -> dict:
    """Test query reasoning and correctness."""
    q = question_data["question"]
    name = f"Query: {q[:50]}..."
    t0 = time.time()
    try:
        result = run_workflow("query", {"project": TEST_PROJECT, "query": q, "top_k": 5})
        duration = round(time.time() - t0, 2)
        success = "error" not in result
        answer = str(result.get("answer", "")).lower()
        # Check for expected keywords
        keyword_hits = [kw for kw in question_data["expected_keywords"] if kw.lower() in answer]
        reasoning_score = len(keyword_hits) / len(question_data["expected_keywords"])
        has_citations = "source" in answer or "citation" in answer or "document" in answer
        log(f"{'PASS' if success else 'FAIL'}: {name} ({duration}s) — reasoning_score={reasoning_score:.0%}, citations={has_citations}")
        if not success:
            log(f"  ERROR: {result.get('error')}")
        return {
            "name": name,
            "question": q,
            "success": success,
            "duration_sec": duration,
            "reasoning_score": reasoning_score,
            "has_citations": has_citations,
            "answer_preview": str(result.get("answer", ""))[:200],
            "error": result.get("error") if not success else None,
        }
    except Exception as e:
        duration = round(time.time() - t0, 2)
        log(f"EXCEPTION: {name} ({duration}s) — {e}")
        return {"name": name, "question": q, "success": False, "duration_sec": duration, "error": str(e)}


def test_rfi(rfi_data: dict) -> dict:
    """Test RFI drafting quality."""
    q = rfi_data["question"]
    name = f"RFI: {q[:50]}..."
    t0 = time.time()
    try:
        result = run_workflow("rfi", {
            "project": TEST_PROJECT,
            "question": q,
            "number": f"RFI-BENCH-{hashlib.md5(q.encode()).hexdigest()[:6].upper()}",
        })
        duration = round(time.time() - t0, 2)
        success = "error" not in result
        rfi_text = str(result.get("rfi_draft", "")).lower()
        has_references = any(ref.lower() in rfi_text for ref in rfi_data["expected_references"])
        has_question = "?" in str(result.get("rfi_draft", ""))
        log(f"{'PASS' if success else 'FAIL'}: {name} ({duration}s) — references={has_references}, question={has_question}")
        if not success:
            log(f"  ERROR: {result.get('error')}")
        return {
            "name": name,
            "question": q,
            "success": success,
            "duration_sec": duration,
            "has_references": has_references,
            "has_question": has_question,
            "rfi_preview": str(result.get("rfi_draft", ""))[:300],
            "error": result.get("error") if not success else None,
        }
    except Exception as e:
        duration = round(time.time() - t0, 2)
        log(f"EXCEPTION: {name} ({duration}s) — {e}")
        return {"name": name, "question": q, "success": False, "duration_sec": duration, "error": str(e)}


def test_scan() -> dict:
    """Test contradiction scanning."""
    name = "Scan: Full project contradiction detection"
    t0 = time.time()
    try:
        result = run_workflow("scan", {"project": TEST_PROJECT})
        duration = round(time.time() - t0, 2)
        success = "error" not in result
        contradictions = result.get("contradictions", [])
        has_critical = any(c.get("severity") == "critical" for c in contradictions)
        has_high = any(c.get("severity") == "high" for c in contradictions)
        log(f"{'PASS' if success else 'FAIL'}: {name} ({duration}s) — found={len(contradictions)}, critical={has_critical}, high={has_high}")
        if not success:
            log(f"  ERROR: {result.get('error')}")
        return {
            "name": name,
            "success": success,
            "duration_sec": duration,
            "contradiction_count": len(contradictions),
            "has_critical": has_critical,
            "has_high": has_high,
            "contradictions_preview": [c.get("description", "")[:100] for c in contradictions[:3]],
            "error": result.get("error") if not success else None,
        }
    except Exception as e:
        duration = round(time.time() - t0, 2)
        log(f"EXCEPTION: {name} ({duration}s) — {e}")
        return {"name": name, "success": False, "duration_sec": duration, "error": str(e)}


def main():
    log("=" * 70)
    log("MEDHA COMPREHENSIVE DEMO BENCHMARK")
    log("=" * 70)
    log(f"Project: {TEST_PROJECT}")
    log(f"Timing Budget: {TIMING_BUDGET_SEC}s")

    overall_t0 = time.time()
    all_results = []

    # Phase 1: Generate sample documents
    generated_docs = generate_sample_documents()
    all_results.append({"phase": "document_generation", "count": len(generated_docs), "types": list(set(d[0] for d in generated_docs))})

    # Phase 2: Ingest all documents
    log("\n" + "=" * 70)
    log("PHASE 2: DOCUMENT INGESTION")
    log("=" * 70)
    ingest_results = []
    for doc_type, file_path in generated_docs:
        ingest_results.append(test_ingest(doc_type, file_path))
    all_results.append({"phase": "ingestion", "tests": ingest_results})

    # Phase 3: Query benchmark
    log("\n" + "=" * 70)
    log("PHASE 3: QUERY REASONING & CORRECTNESS")
    log("=" * 70)
    query_results = []
    for q_data in BENCHMARK_QUESTIONS:
        query_results.append(test_query(q_data))
    all_results.append({"phase": "query", "tests": query_results})

    # Phase 4: RFI benchmark
    log("\n" + "=" * 70)
    log("PHASE 4: RFI DRAFTING QUALITY")
    log("=" * 70)
    rfi_results = []
    for r_data in RFI_TEST_QUESTIONS:
        rfi_results.append(test_rfi(r_data))
    all_results.append({"phase": "rfi", "tests": rfi_results})

    # Phase 5: Contradiction scan
    log("\n" + "=" * 70)
    log("PHASE 5: CONTRADICTION DETECTION")
    log("=" * 70)
    scan_result = test_scan()
    all_results.append({"phase": "scan", "tests": [scan_result]})

    # Summary
    total_duration = round(time.time() - overall_t0, 2)
    all_tests = ingest_results + query_results + rfi_results + [scan_result]
    passed = sum(1 for t in all_tests if t["success"])
    failed = len(all_tests) - passed
    avg_reasoning = sum(r.get("reasoning_score", 0) for r in query_results) / max(len(query_results), 1)
    total_contradictions = scan_result.get("contradiction_count", 0)

    log("\n" + "=" * 70)
    log("BENCHMARK SUMMARY")
    log("=" * 70)
    log(f"Total Duration:     {total_duration}s (budget: {TIMING_BUDGET_SEC}s)")
    log(f"Tests Passed:       {passed}/{len(all_tests)}")
    log(f"Tests Failed:       {failed}/{len(all_tests)}")
    log(f"Avg Reasoning:      {avg_reasoning:.0%}")
    log(f"Contradictions:     {total_contradictions}")
    log(f"Budget Met:         {'YES' if total_duration <= TIMING_BUDGET_SEC else 'NO'}")

    report = {
        "timestamp": datetime.utcnow().isoformat() + "Z",
        "project": TEST_PROJECT,
        "total_duration_sec": total_duration,
        "timing_budget_met": total_duration <= TIMING_BUDGET_SEC,
        "tests_total": len(all_tests),
        "tests_passed": passed,
        "tests_failed": failed,
        "avg_reasoning_score": avg_reasoning,
        "contradictions_found": total_contradictions,
        "document_types_tested": list(set(d[0] for d in generated_docs)),
        "phases": all_results,
        "research_citations": [
            "SaaSFactor2025 — TTFV <10 min or users churn",
            "Li2024 — Multi-agent consensus reduces false negatives by 80%",
            "LayerTeam2025 — RFIs with citations get answered 40% faster",
            "DocumentCrunch2025 — Good RFIs need exact drawing/spec references",
            "Nielsen1994 — System status visibility",
        ],
    }

    REPORT_DIR.mkdir(parents=True, exist_ok=True)
    report_path = REPORT_DIR / f"demo_benchmark_{datetime.now().strftime('%Y%m%d_%H%M%S')}.json"
    with open(report_path, "w") as f:
        json.dump(report, f, indent=2)
    log(f"\nReport saved: {report_path}")

    sys.exit(0 if failed == 0 else 1)


if __name__ == "__main__":
    main()
